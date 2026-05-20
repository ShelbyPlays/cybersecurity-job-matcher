from docx import Document
from docx.shared import Pt, Inches
import os
import pandas as pd


OUTPUT_FOLDER = "cover_letters"


def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)


def add_paragraph(doc, text, bold=False, size=11):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(10)

    run = p.add_run(text)
    run.bold = bold
    run.font.name = "Calibri"
    run.font.size = Pt(size)


def create_safe_filename(text):
    import re
    text = re.sub(r"[^a-zA-Z0-9]+", "_", str(text))
    return text.strip("_")


with open("latest_cover_letter.txt", "r", encoding="utf-8") as f:
    cover_letter = f.read()


doc = Document()
set_margins(doc)

paragraphs = cover_letter.split("\n\n")

for paragraph in paragraphs:
    if paragraph.strip():
        add_paragraph(doc, paragraph.strip())


os.makedirs(OUTPUT_FOLDER, exist_ok=True)

try:
    selected_df = pd.read_csv("selected_job.csv")

    company = selected_df.iloc[0]["Company"]
    title = selected_df.iloc[0]["Title"]

    filename = f"{create_safe_filename(company)}_{create_safe_filename(title)}_Cover_Letter.docx"

except:
    filename = "Cover_Letter.docx"


output_path = os.path.join(OUTPUT_FOLDER, filename)

doc.save(output_path)

print("Cover letter exported.")
print(f"Saved to: {output_path}")