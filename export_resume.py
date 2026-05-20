import json
import os
import re
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


OUTPUT_FOLDER = "resumes"
DEFAULT_OUTPUT_FILE = "tailored_resume.docx"

HEADER_NAME = "SHELBY DEUTSCH"
HEADER_CONTACT = "PHILADELPHIA, PA | (484) 988-0971 | DEUTSCHSHELBY@GMAIL.COM"

EDUCATION = [
    ("M.A. – Cybersecurity", "Roger Williams University — Bristol, RI", "May 2026"),
    ("BS – Cybersecurity; Digital Forensics Minor", "Roger Williams University — Bristol, RI", "May 2025"),
]

CERTIFICATIONS = [
    ("Hybrid Server Pro", "TestOut", "December 2024"),
    ("Pentest+ Pro", "TestOut", "May 2026"),
]

ADDITIONAL_EXPERIENCE = [
    {
        "title": "Online Fitness Coach & Certified Personal Trainer",
        "org": "Solopreneur",
        "date": "2022–Present",
        "bullets": [
            "Manage client coaching operations involving program development, progress tracking, accountability systems, scheduling, and digital communication for individual and group fitness clients.",
            "Develop structured fitness and nutrition programs while maintaining client engagement, performance analysis, and long-term goal planning."
        ]
    }
]


def create_safe_filename(text):
    text = str(text)
    text = re.sub(r"[^a-zA-Z0-9]+", "_", text)
    text = text.strip("_")
    return text[:80]


def set_margins(doc):
    for section in doc.sections:
        section.top_margin = Inches(0.35)
        section.bottom_margin = Inches(0.35)
        section.left_margin = Inches(0.5)
        section.right_margin = Inches(0.5)


def set_base_font(run, size=8.7):
    run.font.name = "Calibri"
    run.font.size = Pt(size)


def add_header_line(doc, text, bold=False, size=9):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1

    run = p.add_run(text)
    run.bold = bold
    set_base_font(run, size)


def add_summary(doc, summary_lines):
    summary_text = " ".join(
        str(line).strip()
        for line in summary_lines
        if str(line).strip()
    )

    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Inches(0.22)
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.line_spacing = 1

    run = p.add_run(summary_text)
    set_base_font(run, 8.7)


def add_section_title(doc, title):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(1)
    p.paragraph_format.line_spacing = 1

    run = p.add_run(title)
    run.bold = True
    set_base_font(run, 9.3)


def add_normal_line(doc, text, bold=False, italic=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1

    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    set_base_font(run, 8.7)


def add_three_part_line(doc, title, location, date):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1

    title_run = p.add_run(title)
    title_run.bold = True
    set_base_font(title_run, 8.7)

    divider1 = p.add_run(" | ")
    set_base_font(divider1, 8.7)

    loc_run = p.add_run(location)
    loc_run.italic = True
    set_base_font(loc_run, 8.7)

    divider2 = p.add_run(" | ")
    set_base_font(divider2, 8.7)

    date_run = p.add_run(date)
    date_run.italic = True
    set_base_font(date_run, 8.7)


def add_bullet(doc, text):
    clean = text.replace("•", "").replace("▪", "").strip()

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.15)
    p.paragraph_format.first_line_indent = Inches(-0.1)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.line_spacing = 1

    run = p.add_run("▪ " + clean)
    set_base_font(run, 8.7)


def add_skills(doc, lines):
    add_section_title(doc, "TECHNICAL SKILLS")

    for line in lines:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.line_spacing = 1

        if ":" in line:
            first, rest = line.split(":", 1)

            bold_run = p.add_run(first + ":")
            bold_run.bold = True
            set_base_font(bold_run, 8.7)

            normal_run = p.add_run(rest)
            set_base_font(normal_run, 8.7)
        else:
            run = p.add_run(line)
            set_base_font(run, 8.7)


def add_education(doc):
    add_section_title(doc, "EDUCATION")
    for title, location, date in EDUCATION:
        add_three_part_line(doc, title, location, date)


def add_certifications(doc):
    add_section_title(doc, "CERTIFICATIONS")
    for title, issuer, date in CERTIFICATIONS:
        add_three_part_line(doc, title, issuer, date)


def add_projects(doc, lines):
    add_section_title(doc, "TECHNICAL PROJECTS")

    bold_titles = [
        "Explainable AI-Driven Anomaly Detection API for Authentication Logs",
        "Password Security Analysis & Awareness Tool",
        "Password Security Analysis & awarness tool",
    ]

    for line in lines:
        if not str(line).strip():
            continue

        clean = str(line).strip()

        if clean.startswith("•") or clean.startswith("▪"):
            add_bullet(doc, clean)
        elif clean in bold_titles:
            add_normal_line(doc, clean, bold=True)
        else:
            add_normal_line(doc, clean)


def add_related_experience(doc, lines):
    add_section_title(doc, "RELATED EXPERIENCE")

    role_titles = [
        "Cybersecurity Graduate Assistant",
        "Cybersecurity Intern",
        "Cybersecurity & Digital Forensic Analysis Intern"
    ]

    i = 0
    while i < len(lines):
        line = str(lines[i]).strip()

        if not line:
            i += 1
            continue

        if line in role_titles:
            title = line
            location = str(lines[i + 1]).strip() if i + 1 < len(lines) else ""
            date = str(lines[i + 2]).strip() if i + 2 < len(lines) else ""

            add_three_part_line(doc, title, location, date)
            i += 3
            continue

        if line.startswith("•") or line.startswith("▪"):
            add_bullet(doc, line)
        else:
            add_normal_line(doc, line)

        i += 1


def add_additional_experience(doc):
    add_section_title(doc, "ADDITIONAL EXPERIENCE")

    for role in ADDITIONAL_EXPERIENCE:
        add_three_part_line(doc, role["title"], role["org"], role["date"])

        for bullet in role["bullets"]:
            add_bullet(doc, bullet)


def main():
    with open("tailored_resume.json", "r", encoding="utf-8") as f:
        tailored = json.load(f)

    doc = Document()
    set_margins(doc)

    add_header_line(doc, HEADER_NAME, bold=True, size=12)
    add_header_line(doc, HEADER_CONTACT, size=8.8)

    add_summary(doc, tailored.get("SUMMARY", []))

    add_skills(doc, tailored.get("TECHNICAL_SKILLS", []))
    add_education(doc)
    add_certifications(doc)
    add_projects(doc, tailored.get("TECHNICAL_PROJECTS", []))
    add_related_experience(doc, tailored.get("RELATED_EXPERIENCE", []))
    add_additional_experience(doc)

    os.makedirs(OUTPUT_FOLDER, exist_ok=True)

    try:
        selected_df = pd.read_csv("selected_job.csv")
        company = str(selected_df.iloc[0]["Company"])
        title = str(selected_df.iloc[0]["Title"])

        safe_company = create_safe_filename(company)
        safe_title = create_safe_filename(title)

        filename = f"{safe_company}_{safe_title}.docx"

    except Exception:
        filename = "Tailored_Resume.docx"

    versioned_output = os.path.join(OUTPUT_FOLDER, filename)

    doc.save(DEFAULT_OUTPUT_FILE)
    doc.save(versioned_output)

    print("Resume export complete.")
    print(f"Created {DEFAULT_OUTPUT_FILE}")
    print(f"Saved versioned copy to {versioned_output}")


if __name__ == "__main__":
    main()